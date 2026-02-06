"""
Resume Parser Module
Extracts text and structured data from various resume formats (PDF, DOCX, TXT)
"""

import io
import re
import logging
from typing import Dict, Optional, List
from pathlib import Path

# PDF parsing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """Parse resumes from various file formats"""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt']
    
    def __init__(self):
        self.file_path = None
        self.file_type = None
        self.raw_text = None
    
    def parse_file(self, file_path: str) -> str:
        """
        Parse resume file and extract text
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.file_path = file_path
        self.file_type = file_path.suffix.lower()
        
        if self.file_type not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {self.file_type}")
        
        logger.info(f"Parsing {self.file_type} file: {file_path}")
        
        if self.file_type == '.pdf':
            self.raw_text = self._parse_pdf(file_path)
        elif self.file_type == '.docx':
            self.raw_text = self._parse_docx(file_path)
        elif self.file_type == '.txt':
            self.raw_text = self._parse_txt(file_path)
        
        logger.info(f"Extracted {len(self.raw_text)} characters from resume")
        return self.raw_text
    
    def parse_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse resume from bytes (for file uploads)
        
        Args:
            file_bytes: File content as bytes
            filename: Name of the file
            
        Returns:
            Extracted text content
        """
        file_type = Path(filename).suffix.lower()
        
        if file_type not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_type}")
        
        self.file_type = file_type
        logger.info(f"Parsing {file_type} from bytes")
        
        if file_type == '.pdf':
            self.raw_text = self._parse_pdf_bytes(file_bytes)
        elif file_type == '.docx':
            self.raw_text = self._parse_docx_bytes(file_bytes)
        elif file_type == '.txt':
            self.raw_text = file_bytes.decode('utf-8', errors='ignore')
        
        return self.raw_text
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                
            return '\n'.join(text)
        
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise
    
    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """Parse PDF from bytes"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        text = []
        
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            
            return '\n'.join(text)
        
        except Exception as e:
            logger.error(f"Error parsing PDF bytes: {e}")
            raise
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise
    
    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        """Parse DOCX from bytes"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        
        except Exception as e:
            logger.error(f"Error parsing DOCX bytes: {e}")
            raise
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise
    
    def extract_structured_data(self, text: str) -> Dict:
        """
        Extract structured information from resume text
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with structured data
        """
        structured_data = {
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'education': self._extract_education(text),
            'experience_years': self._estimate_experience(text)
        }
        
        return structured_data
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name (usually in first few lines)"""
        lines = text.split('\n')
        
        # Name is usually in the first 3 lines
        for line in lines[:3]:
            line = line.strip()
            # Simple heuristic: name is 2-4 words, each capitalized, no numbers
            if line and not re.search(r'\d', line):
                words = line.split()
                if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
                    return line
        
        return None
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group() if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number"""
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{10}',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group()
        
        return None
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education degrees"""
        degrees = []
        degree_patterns = [
            r'(B\.?S\.?|Bachelor of Science|Bachelor\'s)',
            r'(M\.?S\.?|Master of Science|Master\'s)',
            r'(Ph\.?D\.?|Doctorate)',
            r'(B\.?A\.?|Bachelor of Arts)',
            r'(M\.?B\.?A\.?|Master of Business Administration)'
        ]
        
        for pattern in degree_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                degrees.append(re.search(pattern, text, re.IGNORECASE).group())
        
        return degrees
    
    def _estimate_experience(self, text: str) -> Optional[int]:
        """Estimate years of experience"""
        # Look for patterns like "5 years", "5+ years", "5-7 years"
        experience_patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience\s+(?:of\s+)?(\d+)\+?\s*years?'
        ]
        
        for pattern in experience_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        
        return None


def test_resume_parser():
    """Test the resume parser"""
    parser = ResumeParser()
    
    # Test with sample text
    sample_resume = """
    John Doe
    john.doe@email.com | +1-555-123-4567
    
    PROFESSIONAL SUMMARY
    Senior Software Engineer with 5+ years of experience
    
    EDUCATION
    B.S. in Computer Science, MIT, 2018
    M.S. in Computer Science, Stanford, 2020
    
    EXPERIENCE
    Senior Software Engineer, Tech Corp (2020-Present)
    Software Engineer, StartUp Inc (2018-2020)
    """
    
    # Extract structured data
    structured = parser.extract_structured_data(sample_resume)
    print("Structured Data:")
    for key, value in structured.items():
        print(f"  {key}: {value}")


if __name__ == "__main__":
    test_resume_parser()